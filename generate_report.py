import subprocess, os, json, sys, time
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT_ROOT = "/mnt/c/Users/Start/Desktop/1"

def run_cmd(cmd, cwd=None):
    result = subprocess.run(cmd, shell=True, capture_output=True, text=True, cwd=cwd or PROJECT_ROOT, timeout=30)
    return result.stdout + result.stderr

def capture_terminal_outputs():
    results = {}

    # Health check
    print("[1/6] Health check...")
    results["health_check"] = run_cmd("curl -s http://localhost:5000/api/health")

    # Get all notes
    print("[2/6] Getting notes...")
    results["notes_list"] = run_cmd("curl -s http://localhost:5000/api/notes | python3 -m json.tool 2>/dev/null || curl -s http://localhost:5000/api/notes")

    # Tests
    print("[3/6] Running tests...")
    tests = run_cmd(
        "python3 -m pytest tests/ -v --cov=./ --cov-report=term 2>&1",
        cwd=os.path.join(PROJECT_ROOT, "backend")
    )
    results["tests"] = tests

    # Lint
    print("[4/6] Running linter...")
    results["lint"] = run_cmd(
        "python3 -m flake8 . --count --select=E9,F63,F7,F82 --show-source --statistics 2>&1",
        cwd=os.path.join(PROJECT_ROOT, "backend")
    )
    results["lint_all"] = run_cmd(
        "python3 -m flake8 . --count --exit-zero --max-complexity=10 --max-line-length=100 --statistics 2>&1",
        cwd=os.path.join(PROJECT_ROOT, "backend")
    )

    # Tree
    print("[5/6] Tree structure...")
    results["tree"] = run_cmd(
        "find . -maxdepth 4 -not -path '*/node_modules/*' -not -path '*/.git/*' -not -path '*/__pycache__/*' -not -path '*/htmlcov/*' -not -path '*/.pytest_cache/*' -not -path '*/build/*' -not -path '*/venv/*' -not -path '*/Отчет/*' | sort | head -60",
        cwd=PROJECT_ROOT
    )

    # Coverage HTML report exists
    print("[6/6] Checking coverage...")
    results["coverage_exists"] = os.path.exists(os.path.join(PROJECT_ROOT, "backend", "htmlcov", "index.html"))
    results["coverage_content"] = "Coverage HTML report available in backend/htmlcov/" if results["coverage_exists"] else "N/A"

    return results

def create_report(terminal_outputs):
    doc = Document()

    # Page margins: left 30mm, right 10mm, top 20mm, bottom 20mm
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        # Page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        from docx.oxml.ns import qn
        fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._r.append(fldChar1)
        instrText = run._r.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run._r.append(instrText)
        fldChar2 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run._r.append(fldChar2)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5

    # ===== TITLE PAGE =====
    for _ in range(6):
        doc.add_paragraph()

    lines_title = [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        'ФГБОУ ВО «КНИТУ»',
    ]
    for line in lines_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ОТЧЕТ ПО УЧЕБНОЙ ПРАКТИКЕ")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Тема: "Разработка и настройка CI/CD-пайплайна\nдля веб-приложения управления заметками с тегами"')
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    for line in [
        "Направление подготовки: 02.04.03 — Прикладная математика и информатика",
        "Профиль: Технология программирования",
        "Вид практики: Учебная практика",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("2026").font.size = Pt(14)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    h = doc.add_heading('СОДЕРЖАНИЕ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        "Введение",
        "1. Анализ предметной области и постановка задачи",
        "  1.1. Описание предметной области",
        "  1.2. Функциональные требования к приложению",
        "  1.3. Архитектура приложения",
        "2. Подготовка проекта и контейнеризация",
        "  2.1. Структура репозитория",
        "  2.2. Контейнеризация с Docker",
        "  2.3. Docker Compose для оркестрации",
        "3. Реализация CI/CD-пайплайна",
        "  3.1. Непрерывная интеграция (CI)",
        "  3.2. Непрерывное развертывание (CD)",
        "  3.3. Защита ветки main",
        "4. Тестирование пайплайна",
        "  4.1. Локальное тестирование",
        "  4.2. Результаты выполнения CI/CD в GitHub Actions",
        "5. Инструкция по запуску проекта",
        "  5.1. Запуск через Docker",
        "  5.2. Запуск backend локально",
        "  5.3. Запуск frontend локально",
        "  5.4. Запуск тестов и линтеров",
        "  5.5. Ручная проверка API",
        "Заключение",
        "Список использованных источников",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(14)

    doc.add_page_break()

    # ===== INTRODUCTION =====
    h = doc.add_heading('ВВЕДЕНИЕ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'Целью учебной практики является освоение современных инженерных практик непрерывной '
        'интеграции и непрерывного развертывания (CI/CD) при разработке программных систем.'
    )
    doc.add_paragraph(
        'В ходе практики были решены следующие задачи:'
    )
    tasks = [
        'освоение работы с системой контроля версий Git;',
        'реализация автоматизированного CI-pipeline;',
        'реализация автоматического развертывания приложения (CD);',
        'обеспечение воспроизводимости и управляемости окружений с помощью Docker;',
        'документирование инженерных решений.',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph(
        'В качестве индивидуального проекта разработано веб-приложение управления заметками с тегами, '
        'состоящее из backend на Python Flask и frontend на React.'
    )

    # ===== CHAPTER 1 =====
    doc.add_page_break()
    h = doc.add_heading('1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПОСТАНОВКА ЗАДАЧИ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1.1. Описание предметной области', level=2)
    doc.add_paragraph(
        'Предметная область — веб-приложение для создания и управления заметками. '
        'Пользователь может создавать, просматривать, редактировать и удалять заметки, '
        'а также назначать им теги для удобной категоризации и поиска.'
    )

    doc.add_heading('1.2. Функциональные требования', level=2)
    func_reqs = [
        'CRUD-операции для заметок (создание, чтение, обновление, удаление);',
        'Назначение тегов заметкам;',
        'Поиск заметок по тегам;',
        'REST API для взаимодействия с backend;',
        'Пользовательский веб-интерфейс на React.',
    ]
    for r in func_reqs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('1.3. Архитектура приложения', level=2)
    doc.add_paragraph(
        'Приложение построено по архитектуре клиент-сервер:'
    )
    arch_items = [
        'Backend: Python Flask (REST API) + SQLite (база данных)',
        'Frontend: React 18 + Tailwind CSS + Axios',
        'Контейнеризация: Docker + Docker Compose',
        'CI/CD: GitHub Actions',
    ]
    for a in arch_items:
        doc.add_paragraph(a, style='List Bullet')

    # API endpoints table
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Таблица 1 — Эндпоинты REST API')
    run.bold = True

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Метод', 'Эндпоинт', 'Описание']
    data = [
        ['GET', '/api/health', 'Проверка работоспособности API'],
        ['GET', '/api/notes', 'Получить все заметки'],
        ['GET', '/api/notes/<id>', 'Получить заметку по ID'],
        ['POST', '/api/notes', 'Создать новую заметку'],
        ['PUT', '/api/notes/<id>', 'Обновить заметку'],
        ['DELETE', '/api/notes/<id>', 'Удалить заметку'],
        ['GET', '/api/notes/tag/<tag>', 'Поиск заметок по тегу'],
    ]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_text

    # Model code
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 1 — Модель данных Note (CRUD операции)')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'backend', 'models.py')) as f:
        model_code = f.read()
    p = doc.add_paragraph()
    run = p.add_run(model_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # API code
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 2 — Основные эндпоинты API')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'backend', 'app.py')) as f:
        api_code = f.read()
    p = doc.add_paragraph()
    run = p.add_run(api_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # ===== CHAPTER 2 =====
    doc.add_page_break()
    h = doc.add_heading('2. ПОДГОТОВКА ПРОЕКТА И КОНТЕЙНЕРИЗАЦИЯ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.1. Структура репозитория', level=2)
    doc.add_paragraph(
        'Проект представляет собой монорепозиторий, содержащий backend и frontend в одном репозитории.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Листинг 3 — Структура репозитория проекта')
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(terminal_outputs.get('tree', 'Не получено'))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_paragraph(
        'Рисунок 1 — Структура репозитория в GitHub\n'
        '(требуется скриншот страницы репозитория на GitHub или вывод команды tree)'
    )

    doc.add_heading('2.2. Контейнеризация с Docker', level=2)
    doc.add_paragraph(
        'Для каждого компонента приложения создан Dockerfile, обеспечивающий '
        'воспроизводимость окружения. Backend использует официальный образ Python 3.9-slim, '
        'frontend — multi-stage сборку: сборка в node:18-alpine, запуск через nginx:alpine.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Листинг 4 — Dockerfile для backend')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'backend', 'Dockerfile')) as f:
        docker_backend = f.read()
    p = doc.add_paragraph()
    run = p.add_run(docker_backend)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Листинг 5 — Dockerfile для frontend (multi-stage)')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'frontend', 'Dockerfile')) as f:
        docker_frontend = f.read()
    p = doc.add_paragraph()
    run = p.add_run(docker_frontend)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 6 — Конфигурация nginx (прокси на backend)')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'frontend', 'nginx.conf')) as f:
        nginx_conf = f.read()
    p = doc.add_paragraph()
    run = p.add_run(nginx_conf)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('2.3. Docker Compose для оркестрации', level=2)
    doc.add_paragraph(
        'Для одновременного запуска всех сервисов используется Docker Compose. '
        'Файл docker-compose.yml описывает два сервиса: backend и frontend, '
        'общую сеть (notes-network) и том для базы данных (notes-database). '
        'Backend содержит healthcheck для проверки доступности API, '
        'frontend зависит от backend с условием service_healthy.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Листинг 7 — docker-compose.yml')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, 'docker-compose.yml')) as f:
        compose = f.read()
    p = doc.add_paragraph()
    run = p.add_run(compose)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # ===== CHAPTER 3 =====
    doc.add_page_break()
    h = doc.add_heading('3. РЕАЛИЗАЦИЯ CI/CD-ПАЙПЛАЙНА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.1. Непрерывная интеграция (CI)', level=2)
    doc.add_paragraph(
        'CI-пайплайн настроен в файле .github/workflows/ci.yml и запускается '
        'при каждом push и Pull Request в ветки main и develop.'
    )
    doc.add_paragraph('Pipeline состоит из следующих этапов:')
    ci_steps = [
        'Backend — Lint & Test: проверка стиля (flake8), запуск unit-тестов (pytest) с покрытием, проверка импорта Flask-приложения.',
        'Frontend — Lint & Test: проверка стиля (ESLint), сборка React-приложения, запуск тестов.',
        'Docker Build: сборка Docker-образов backend и frontend, валидация docker-compose конфигурации. Выполняется только после успешного прохождения первых двух этапов.',
        'Integration Validation: проверка наличия всех обязательных файлов проекта, вывод статистики.',
    ]
    for step in ci_steps:
        doc.add_paragraph(step, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Листинг 8 — CI pipeline (GitHub Actions — .github/workflows/ci.yml)')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, '.github', 'workflows', 'ci.yml')) as f:
        ci_yml = f.read()
    p = doc.add_paragraph()
    run = p.add_run(ci_yml)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('3.2. Непрерывное развертывание (CD)', level=2)
    doc.add_paragraph(
        'CD-пайплайн настроен в файле .github/workflows/cd.yml и запускается '
        'автоматически при push в ветку main, а также может быть запущен вручную через workflow_dispatch.'
    )
    doc.add_paragraph('Этапы CD-пайплайна:')
    cd_steps = [
        'Сборка Docker-образов backend и frontend с тегом коммита (github.sha);',
        'Локальное тестирование деплоя: запуск контейнеров, проверка здоровья backend, проверка доступности frontend;',
        'Генерация отчета о деплое в формате Markdown;',
        'Сохранение отчета как артефакта workflow (доступен для скачивания 30 дней);',
        'Уведомление об успешном или ошибочном выполнении.',
    ]
    for step in cd_steps:
        doc.add_paragraph(step, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Листинг 9 — CD pipeline (GitHub Actions — .github/workflows/cd.yml)')
    run.bold = True

    with open(os.path.join(PROJECT_ROOT, '.github', 'workflows', 'cd.yml')) as f:
        cd_yml = f.read()
    p = doc.add_paragraph()
    run = p.add_run(cd_yml)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('3.3. Защита ветки main', level=2)
    doc.add_paragraph(
        'Для ветки main настроены правила защиты в GitHub (Settings → Branches → Add branch protection rule):'
    )
    protection = [
        'Require a pull request before merging — требуется создание Pull Request перед слиянием;',
        'Require status checks to pass before merging — требуется успешное прохождение всех CI-проверок;',
        'В качестве обязательных статус-чеков выбраны: Backend - Lint & Test, Frontend - Lint & Test, Docker Build, Integration Validation.',
    ]
    for p_text in protection:
        doc.add_paragraph(p_text, style='List Bullet')
    doc.add_paragraph(
        'Это гарантирует, что в ветку main попадает только код, успешно прошедший все проверки.'
    )
    doc.add_paragraph(
        'Рисунок 2 — Правила защиты ветки main\n'
        '(требуется скриншот Settings → Branches → Branch protection rules вашего репозитория)'
    )

    # ===== CHAPTER 4 =====
    doc.add_page_break()
    h = doc.add_heading('4. ТЕСТИРОВАНИЕ ПАЙПЛАЙНА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.1. Локальное тестирование', level=2)

    # Health check
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 10 — Проверка работоспособности API')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(terminal_outputs.get('health_check', 'Не получено'))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Notes list
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 11 — Получение списка заметок через API')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(terminal_outputs.get('notes_list', 'Не получено'))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Tests
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 12 — Результаты запуска unit-тестов с покрытием')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(terminal_outputs.get('tests', 'Не получено'))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Lint
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Листинг 13 — Результаты проверки стиля кода (flake8)')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(terminal_outputs.get('lint', 'Не получено'))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Coverage
    if terminal_outputs.get("coverage_exists"):
        doc.add_paragraph(
            'Отчет о покрытии кода тестами доступен в формате HTML в директории backend/htmlcov/. '
            'Для просмотра откройте файл index.html в браузере.'
        )
        doc.add_paragraph(
            'Рисунок 3 — Отчет о покрытии кода тестами\n'
            '(требуется скриншот файла backend/htmlcov/index.html, открытого в браузере)'
        )

    # App screenshots instructions
    doc.add_paragraph()
    doc.add_paragraph(
        'Для визуального подтверждения работы приложения необходимо сделать следующие скриншоты:'
    )
    screenshot_instructions = [
        'Рисунок 4 — Веб-интерфейс приложения управления заметками (http://localhost:3000 после запуска docker-compose up --build). На скриншоте должны быть видны созданные заметки, теги и панель статистики.',
        'Рисунок 5 — Форма создания/редактирования заметки (модальное окно с полями заголовка, содержания и тегов).',
        'Рисунок 6 — Создание заметки через API терминалом (команда curl -X POST ... с ответом сервера).',
        'Рисунок 7 — Проверка здоровья API через браузер или curl (ответ {"status": "healthy", "database": "connected"}).',
    ]
    for s in screenshot_instructions:
        doc.add_paragraph(s)

    doc.add_heading('4.2. Результаты выполнения CI/CD в GitHub Actions', level=2)
    doc.add_paragraph(
        'Для проверки полного цикла CI/CD необходимо:'
    )
    ci_check_steps = [
        'Создать новую ветку и внести изменения в код;',
        'Открыть Pull Request в ветку main;',
        'Дождаться автоматического запуска CI-пайплайна (Actions → CI Pipeline);',
        'Убедиться, что все 4 статус-чека пройдены успешно (зеленые галочки);',
        'Выполнить Merge Pull Request;',
        'Убедиться, что CD-пайплайн запустился автоматически;',
        'Проверить доступность приложения по публичному URL.',
    ]
    for step in ci_check_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph(
        'Необходимые скриншоты из GitHub Actions:'
    )
    gh_screenshots = [
        'Рисунок 8 — Успешное выполнение CI pipeline (страница Actions → CI Pipeline, все jobs — зеленые).',
        'Рисунок 9 — Детальный просмотр этапа Backend - Lint & Test внутри CI pipeline.',
        'Рисунок 10 — Успешное выполнение CD pipeline (страница Actions → CD Pipeline, все шаги пройдены).',
        'Рисунок 11 — Pull Request с пройденными CI-проверками (зеленые галочки внизу PR).',
        'Рисунок 12 — Страница репозитория с бейджами CI/CD статуса в README.md (опционально).',
    ]
    for s in gh_screenshots:
        doc.add_paragraph(s)

    # ===== CHAPTER 5 =====
    doc.add_page_break()
    h = doc.add_heading('5. ИНСТРУКЦИЯ ПО ЗАПУСКУ ПРОЕКТА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('5.1. Запуск через Docker (рекомендованный способ)', level=2)
    doc.add_paragraph('Требования: установленные Docker и Docker Compose.')
    docker_cmds = '''# Клонирование репозитория
git clone <URL вашего репозитория>
cd 1

# Запуск всех сервисов
docker-compose up --build

# После сборки:
# Frontend: http://localhost:3000
# Backend API: http://localhost:5000/api/health

# Остановка сервисов
docker-compose down'''
    p = doc.add_paragraph()
    run = p.add_run(docker_cmds)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('5.2. Запуск backend локально', level=2)
    doc.add_paragraph('Требования: Python 3.9+ установленный.')
    backend_cmds = '''# Перейти в директорию backend
cd backend

# Создать и активировать виртуальное окружение
python -m venv venv
source venv/bin/activate  # Linux/Mac
# venv\\Scripts\\activate  # Windows

# Установить зависимости
pip install -r requirements.txt

# Запустить приложение
python app.py

# Сервер будет доступен на http://localhost:5000'''
    p = doc.add_paragraph()
    run = p.add_run(backend_cmds)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('5.3. Запуск frontend локально', level=2)
    doc.add_paragraph('Требования: Node.js 18+ установленный.')
    frontend_cmds = '''# Перейти в директорию frontend
cd frontend

# Установить зависимости
npm install --legacy-peer-deps

# Запустить в режиме разработки
npm start

# Приложение будет доступно на http://localhost:3000

# Сборка для продакшна
npm run build'''
    p = doc.add_paragraph()
    run = p.add_run(frontend_cmds)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('5.4. Запуск тестов и линтеров', level=2)
    test_cmds = '''# === Backend тесты ===
cd backend

# Unit-тесты с покрытием
pytest tests/ -v --cov=./ --cov-report=term

# Просмотр HTML отчета о покрытии
# Откройте backend/htmlcov/index.html в браузере

# Запуск линтера flake8
flake8 . --count --select=E9,F63,F7,F82 --show-source --statistics
flake8 . --count --exit-zero --max-complexity=10 --max-line-length=100 --statistics


# === Frontend тесты ===
cd frontend

# Запуск тестов React
npm test -- --watchAll=false --passWithNoTests

# Запуск линтера ESLint
npm run lint'''
    p = doc.add_paragraph()
    run = p.add_run(test_cmds)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('5.5. Ручная проверка API', level=2)
    doc.add_paragraph('Примеры запросов к API через curl:')
    api_cmds = '''# Проверка здоровья API
curl http://localhost:5000/api/health

# Создание заметки
curl -X POST http://localhost:5000/api/notes \
  -H "Content-Type: application/json" \
  -d '{"title":"Моя заметка","content":"Текст заметки","tags":["работа","важное"]}'

# Получение всех заметок
curl http://localhost:5000/api/notes

# Получение заметки по ID
curl http://localhost:5000/api/notes/1

# Поиск по тегу
curl http://localhost:5000/api/notes/tag/работа

# Обновление заметки
curl -X PUT http://localhost:5000/api/notes/1 \
  -H "Content-Type: application/json" \
  -d '{"title":"Новый заголовок","content":"Новый текст","tags":["обновлено"]}'

# Удаление заметки
curl -X DELETE http://localhost:5000/api/notes/1'''
    p = doc.add_paragraph()
    run = p.add_run(api_cmds)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # ===== CONCLUSION =====
    doc.add_page_break()
    h = doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'В ходе учебной практики был разработан и настроен полный CI/CD-пайплайн '
        'для веб-приложения управления заметками с тегами.'
    )

    doc.add_paragraph('Результаты выполнения практики:')
    conclusion_items = [
        'Разработано веб-приложение на Flask (backend) и React (frontend) с CRUD-операциями и поиском по тегам.',
        'Созданы Dockerfile для backend и frontend, а также docker-compose.yml для оркестрации контейнеров.',
        'Настроен CI-пайплайн в GitHub Actions, выполняющий линтинг, тестирование и сборку Docker-образов при каждом Pull Request.',
        'Настроен CD-пайплайн для автоматического развертывания и тестирования приложения после слияния в main.',
        'Настроены правила защиты ветки main, требующие успешного прохождения CI перед слиянием.',
        'Подготовлена документация по запуску и эксплуатации проекта.',
    ]
    for item in conclusion_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Все поставленные задачи учебной практики выполнены в полном объеме. '
        'CI/CD-пайплайн обеспечивает автоматизацию проверки качества кода, сборки и развертывания, '
        'что соответствует современным инженерным практикам разработки программного обеспечения.'
    )

    # ===== REFERENCES =====
    doc.add_page_break()
    h = doc.add_heading('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    refs = [
        '1. GitHub Docs. GitHub Actions — официальная документация. — URL: https://docs.github.com/en/actions (дата обращения: 01.07.2026).',
        '2. Docker Documentation. Dockerfile reference. — URL: https://docs.docker.com/engine/reference/builder/ (дата обращения: 01.07.2026).',
        '3. Docker Documentation. Compose file reference. — URL: https://docs.docker.com/compose/compose-file/ (дата обращения: 01.07.2026).',
        '4. Flask Documentation. — URL: https://flask.palletsprojects.com/ (дата обращения: 01.07.2026).',
        '5. React Documentation. — URL: https://react.dev/ (дата обращения: 01.07.2026).',
        '6. pytest Documentation. — URL: https://docs.pytest.org/ (дата обращения: 01.07.2026).',
        '7. Flake8 Documentation. — URL: https://flake8.pycqa.org/ (дата обращения: 01.07.2026).',
        '8. Tailwind CSS Documentation. — URL: https://tailwindcss.com/docs (дата обращения: 01.07.2026).',
        '9. Axios Documentation. — URL: https://axios-http.com/ (дата обращения: 01.07.2026).',
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Save
    output_path = os.path.join(PROJECT_ROOT, 'Отчет', 'Отчет_по_учебной_практике.docx')
    doc.save(output_path)
    print(f"\n=== REPORT SAVED: {output_path} ===")
    return output_path

if __name__ == '__main__':
    print("=== Capturing terminal outputs ===")
    outputs = capture_terminal_outputs()
    print("\n=== Generating report ===")
    report_path = create_report(outputs)
    print(f"\nDone! Report: {report_path}")
